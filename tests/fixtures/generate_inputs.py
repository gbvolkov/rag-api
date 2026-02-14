from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent / "inputs"


def _build_long_paragraph(seed: int) -> str:
    return (
        f"Section {seed}: The knowledge graph subsystem tracks provenance, confidence, and semantic lineage for "
        f"retrieval-augmented generation workflows. It stores typed entities, relations, and temporal snapshots while "
        f"supporting replay-safe updates and deterministic ingestion checkpoints for auditability."
    )


def write_text() -> None:
    paragraphs = [_build_long_paragraph(i) for i in range(1, 81)]
    body = "\n\n".join(paragraphs)
    (ROOT / "long_text.txt").write_text(body, encoding="utf-8")


def write_markdown() -> None:
    lines = ["# Knowledge Base Operations Handbook", ""]
    for idx in range(1, 26):
        lines.append(f"## Workflow {idx}")
        lines.append(_build_long_paragraph(idx))
        lines.append("")
    lines.extend(
        [
            "## Capacity Snapshot",
            "",
            "| Region | Cluster | Nodes | Avg Latency (ms) | Errors/day |",
            "|---|---|---:|---:|---:|",
        ]
    )
    for region in ["us-east", "us-west", "eu-central", "ap-south"]:
        for cluster_idx in range(1, 7):
            lines.append(
                f"| {region} | {region}-c{cluster_idx} | {80 + cluster_idx * 5} | "
                f"{12 + cluster_idx} | {cluster_idx * 3} |"
            )
    (ROOT / "long_markdown.md").write_text("\n".join(lines), encoding="utf-8")


def write_json() -> None:
    items = []
    for idx in range(1, 181):
        items.append(
            {
                "id": idx,
                "title": f"Incident Report {idx}",
                "category": "operations" if idx % 2 == 0 else "retrieval",
                "priority": "high" if idx % 5 == 0 else "normal",
                "summary": _build_long_paragraph(idx),
                "tags": [f"tag-{idx % 7}", f"cluster-{idx % 11}", "kbman"],
                "metrics": {
                    "duration_minutes": (idx % 90) + 10,
                    "failed_tasks": idx % 13,
                    "retries": idx % 4,
                },
                "history": [
                    {"event": "created", "at": f"2025-09-{(idx % 28) + 1:02d}"},
                    {"event": "triaged", "at": f"2025-10-{(idx % 28) + 1:02d}"},
                    {"event": "resolved", "at": f"2025-11-{(idx % 28) + 1:02d}"},
                ],
            }
        )
    payload = {
        "dataset": "kbman_long_fixture",
        "version": 1,
        "notes": "Large nested JSON payload for integration and chunking tests.",
        "items": items,
    }
    (ROOT / "long_data.json").write_text(json.dumps(payload, indent=2), encoding="utf-8")


def write_qa() -> None:
    rows = []
    for idx in range(1, 121):
        rows.append(f"Q: How does retrieval policy {idx} mitigate stale context?")
        rows.append(
            "A: It applies source freshness scoring, lineage-aware deduplication, and score-calibrated merging "
            f"before final ranking. Reference profile {idx} also enables stable pagination checkpoints."
        )
        rows.append("")
    (ROOT / "long_qa.txt").write_text("\n".join(rows), encoding="utf-8")


def write_csv() -> None:
    header = ["id", "title", "region", "category", "owner", "priority", "status", "summary"]
    with (ROOT / "long_data.csv").open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(header)
        for idx in range(1, 361):
            writer.writerow(
                [
                    idx,
                    f"Task {idx}",
                    ["us-east", "us-west", "eu-central", "ap-south"][idx % 4],
                    "retrieval" if idx % 2 else "indexing",
                    f"owner-{idx % 15}",
                    ["low", "normal", "high"][idx % 3],
                    ["open", "running", "done"][idx % 3],
                    _build_long_paragraph(idx),
                ]
            )


def write_table_csv() -> None:
    header = ["region", "product", "quarter", "revenue", "cost", "margin", "owner"]
    with (ROOT / "long_table.csv").open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(header)
        for idx in range(1, 321):
            revenue = 20000 + (idx * 73)
            cost = 9000 + (idx * 41)
            margin = revenue - cost
            writer.writerow(
                [
                    ["us-east", "us-west", "eu-central", "ap-south"][idx % 4],
                    f"product-{idx % 9}",
                    f"2025-Q{(idx % 4) + 1}",
                    revenue,
                    cost,
                    margin,
                    f"owner-{idx % 12}",
                ]
            )


def write_docx() -> None:
    doc = Document()
    doc.add_heading("Knowledge Base Program Manual", 0)
    for section in range(1, 18):
        doc.add_heading(f"Section {section}: Operating Model", level=1)
        for paragraph_idx in range(1, 5):
            doc.add_paragraph(
                f"{_build_long_paragraph(section * paragraph_idx)} "
                f"This paragraph captures implementation notes, rollout constraints, and fallback controls."
            )
        doc.add_heading(f"Section {section}.1: Change Control", level=2)
        doc.add_paragraph(
            "Change windows are grouped by project scope and enforced via immutable version sets for segments and chunks."
        )
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Team"
    hdr[1].text = "Project"
    hdr[2].text = "Artifacts/day"
    hdr[3].text = "Avg latency"
    hdr[4].text = "SLA"
    for idx in range(1, 41):
        row = table.add_row().cells
        row[0].text = f"Team-{idx % 8}"
        row[1].text = f"Project-{idx % 19}"
        row[2].text = str(100 + idx * 3)
        row[3].text = f"{10 + (idx % 9)}ms"
        row[4].text = "99.9%"
    doc.save(ROOT / "long_document.docx")


def write_excel() -> None:
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "events"
    ws1.append(["event_id", "project", "stage", "duration_ms", "owner", "notes"])
    for idx in range(1, 251):
        ws1.append(
            [
                idx,
                f"project-{idx % 24}",
                ["ingest", "segment", "chunk", "index", "retrieve"][idx % 5],
                25 + (idx % 140),
                f"owner-{idx % 11}",
                _build_long_paragraph(idx),
            ]
        )

    ws2 = wb.create_sheet("metrics")
    ws2.append(["date", "requests", "errors", "p50_ms", "p95_ms", "p99_ms"])
    for day in range(1, 121):
        ws2.append(
            [
                f"2025-11-{(day % 28) + 1:02d}",
                1000 + day * 9,
                day % 21,
                12 + (day % 7),
                19 + (day % 10),
                28 + (day % 13),
            ]
        )

    ws3 = wb.create_sheet("taxonomy")
    ws3.append(["domain", "term", "definition"])
    for idx in range(1, 101):
        ws3.append(["retrieval", f"term-{idx}", _build_long_paragraph(idx)])

    wb.save(ROOT / "long_workbook.xlsx")


def write_pdf() -> None:
    path = ROOT / "long_tables.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()
    elements = [Paragraph("KBMan Long PDF Fixture", styles["Heading1"]), Spacer(1, 12)]
    elements.append(Paragraph(_build_long_paragraph(1), styles["BodyText"]))
    elements.append(Spacer(1, 12))

    data = [["Region", "Cluster", "Requests", "Errors", "P95 (ms)"]]
    for idx in range(1, 61):
        data.append(
            [
                ["us-east", "us-west", "eu-central", "ap-south"][idx % 4],
                f"cluster-{idx}",
                str(1200 + idx * 17),
                str(idx % 13),
                str(20 + (idx % 11)),
            ]
        )
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2d4059")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
            ]
        )
    )
    elements.append(table)
    doc.build(elements)


def main() -> None:
    ROOT.mkdir(parents=True, exist_ok=True)
    write_text()
    write_markdown()
    write_json()
    write_qa()
    write_csv()
    write_table_csv()
    write_docx()
    write_excel()
    write_pdf()
    print(f"Generated long input fixtures in: {ROOT}")


if __name__ == "__main__":
    main()
